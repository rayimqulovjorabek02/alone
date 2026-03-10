"""
modules/export.py — Chat eksport (TXT, MD, DOCX, PDF)
"""
import io
from datetime import datetime


def export_chat_txt(messages: list, username: str = "Foydalanuvchi") -> str:
    lines = [f"Alone AI — Suhbat", f"Eksport: {datetime.now().strftime('%d.%m.%Y %H:%M')}", "="*50, ""]
    for msg in messages:
        role = "Siz" if msg.get("role") == "user" else "Alone AI"
        lines.append(f"[{role}]")
        lines.append(msg.get("content", ""))
        lines.append("")
    return "\n".join(lines)


def export_chat_md(messages: list, username: str = "Foydalanuvchi") -> str:
    lines = [f"# Alone AI — Suhbat", f"*Eksport: {datetime.now().strftime('%d.%m.%Y %H:%M')}*", "---", ""]
    for msg in messages:
        role = f"**{username}**" if msg.get("role") == "user" else "**Alone AI**"
        lines.append(f"### {role}")
        lines.append(msg.get("content", ""))
        lines.append("---")
    return "\n".join(lines)


def export_chat_docx(messages: list, username: str = "Foydalanuvchi") -> bytes | None:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        doc = Document()
        doc.add_heading("Alone AI — Suhbat", 0)
        doc.add_paragraph(f"Eksport: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        for msg in messages:
            role = username if msg.get("role") == "user" else "Alone AI"
            p = doc.add_paragraph()
            run = p.add_run(f"{role}: ")
            run.bold = True
            p.add_run(msg.get("content", ""))
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return None


def export_chat_pdf(messages: list, username: str = "Foydalanuvchi") -> bytes | None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = [Paragraph("Alone AI — Suhbat", styles["Title"]), Spacer(1, 12)]
        for msg in messages:
            role = username if msg.get("role") == "user" else "Alone AI"
            content = msg.get("content", "").replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            story.append(Paragraph(f"<b>{role}:</b> {content}", styles["Normal"]))
            story.append(Spacer(1, 6))
        doc.build(story)
        return buf.getvalue()
    except Exception:
        return None
